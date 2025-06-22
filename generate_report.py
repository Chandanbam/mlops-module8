#!/usr/bin/env python3
"""
Generate Performance Analysis Report as Word Document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime

def create_performance_report():
    """Create a Word document with performance analysis report."""
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Performance Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Timestamp
    timestamp_para = doc.add_paragraph()
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp_para.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp_run.bold = True
    
    doc.add_paragraph()  # Add space
    
    # 1. Pipeline Performance Metrics
    doc.add_heading('1. Pipeline Performance Metrics', level=1)
    
    # Create table
    table1 = doc.add_table(rows=8, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Metric', 'Traditional Pipeline', 'Dask Pipeline']
    for i, header in enumerate(headers):
        cell = table1.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data
    data1 = [
        ['Time (s)', '11.82', '15.13'],
        ['Memory Initial (MB)', '608.13', '698.24'],
        ['Memory Final (MB)', '730.14', '743.88'],
        ['Memory Increase (MB)', '122.01', '45.63'],
        ['CPU Usage (%)', '2.8', '2.1'],
        ['Dataset Size', '206,400', '206,400'],
        ['# Features', '8', '8']
    ]
    
    for i, row_data in enumerate(data1, 1):
        for j, value in enumerate(row_data):
            table1.cell(i, j).text = value
    
    doc.add_paragraph()  # Add space
    
    # 2. Model Performance Comparison
    doc.add_heading('2. Model Performance Comparison', level=1)
    
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers2 = ['Metric', 'Traditional', 'Dask', 'Improvement (%)']
    for i, header in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data
    data2 = [
        ['MSE', '0.2004', '2.3308', '-1063.00'],
        ['RMSE', '0.4477', '1.5267', '-241.03'],
        ['R² Score', '0.8492', '0.8434', '-0.68']
    ]
    
    for i, row_data in enumerate(data2, 1):
        for j, value in enumerate(row_data):
            table2.cell(i, j).text = value
    
    doc.add_paragraph()  # Add space
    
    # 3. Resource Usage Comparison
    doc.add_heading('3. Resource Usage Comparison', level=1)
    
    table3 = doc.add_table(rows=3, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers3 = ['Metric', 'Traditional', 'Dask', 'Improvement (%)']
    for i, header in enumerate(headers3):
        cell = table3.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data
    data3 = [
        ['Training Time (s)', '11.82', '15.13', '-28.03'],
        ['Memory Usage (MB)', '122.01', '45.63', '+62.60']
    ]
    
    for i, row_data in enumerate(data3, 1):
        for j, value in enumerate(row_data):
            table3.cell(i, j).text = value
    
    doc.add_paragraph()  # Add space
    
    # 4. Summary
    doc.add_heading('4. Summary', level=1)
    
    summary_items = [
        'Dask Faster: No',
        'Dask More Memory Efficient: Yes',
        'Dask Scalable: Yes',
        'Overall Winner: Dask (due to memory efficiency and scalability)'
    ]
    
    for item in summary_items:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(item)
    
    doc.add_paragraph()  # Add space
    
    # 5. Configuration Used
    doc.add_heading('5. Configuration Used', level=1)
    
    config_items = [
        'Dataset: California Housing, scale factor 10',
        'Dask: 4 workers, 2 threads/worker, 2GB memory/worker',
        'Model: SGDRegressor, 100 estimators, max depth 10',
        'Training: 5-fold CV, scoring: neg_mean_squared_error',
        'Performance: Profiling, memory tracking, timing analysis enabled'
    ]
    
    for item in config_items:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(item)
    
    doc.add_paragraph()  # Add space
    
    # 6. Performance Comparison Plot
    doc.add_heading('6. Performance Comparison Plot', level=1)
    
    # Add the performance comparison image if it exists
    image_path = 'results/plots/performance_comparison.png'
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6))
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('Performance comparison plot not found.')
    
    # Save the document
    output_path = 'results/reports/performance_report.docx'
    doc.save(output_path)
    print(f"Performance report saved to: {output_path}")

if __name__ == "__main__":
    create_performance_report() 